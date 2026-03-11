from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import Response
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from backend.database import get_db
from backend.models import Project, Chapter
from backend.services import pdf_service
import re

router = APIRouter()


def _safe_name(title: str, ext: str) -> str:
    safe = re.sub(r'[^\w\-_]', '_', title)[:50]
    return f"{safe}.{ext}"


async def _get_project_chapters(project_id: int, db: AsyncSession):
    project = await db.get(Project, project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Roman introuvable")
    result = await db.execute(
        select(Chapter)
        .where(Chapter.project_id == project_id, Chapter.archived == False)
        .order_by(Chapter.order_index)
    )
    return project, result.scalars().all()


# ── PDF ────────────────────────────────────────────────────────────────────────

@router.post("/pdf")
async def export_pdf(body: dict, db: AsyncSession = Depends(get_db)):
    project_id = body.get("project_id")
    if not project_id:
        raise HTTPException(status_code=400, detail="project_id requis")
    project, chapters = await _get_project_chapters(project_id, db)
    try:
        pdf_bytes = pdf_service.generate_novel_pdf(
            {"title": project.title, "subtitle": project.subtitle,
             "author": project.author, "cover_color": project.cover_color},
            [{"title": c.title, "content": c.content} for c in chapters]
        )
        filename = _safe_name(project.title, "pdf")
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erreur PDF : {str(e)}")


# ── Word (.docx) ───────────────────────────────────────────────────────────────

@router.post("/docx")
async def export_docx(body: dict, db: AsyncSession = Depends(get_db)):
    project_id = body.get("project_id")
    if not project_id:
        raise HTTPException(status_code=400, detail="project_id requis")
    project, chapters = await _get_project_chapters(project_id, db)

    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io, re as _re

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Georgia'
        style.font.size = Pt(12)

        doc.add_paragraph('')
        doc.add_paragraph('')
        title_p = doc.add_paragraph(project.title)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.runs[0].font.size = Pt(28)
        title_p.runs[0].font.bold = True

        if project.subtitle:
            sub_p = doc.add_paragraph(project.subtitle)
            sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_p.runs[0].font.size = Pt(14)
            sub_p.runs[0].font.italic = True

        if project.author:
            doc.add_paragraph('')
            auth_p = doc.add_paragraph(project.author)
            auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            auth_p.runs[0].font.size = Pt(13)

        doc.add_page_break()

        def html_to_plain(html):
            h = _re.sub(r'<br\s*/?>', '\n', html)
            h = _re.sub(r'</p>', '\n', h)
            h = _re.sub(r'<[^>]+>', '', h)
            return h.strip()

        for i, chapter in enumerate(chapters):
            doc.add_heading(f"{i+1}. {chapter.title}", level=1)
            if chapter.subtitle:
                sub_p = doc.add_paragraph(chapter.subtitle)
                sub_p.runs[0].font.italic = True
            text = html_to_plain(chapter.content or "")
            for para in text.split('\n'):
                if para.strip():
                    doc.add_paragraph(para.strip())
            doc.add_page_break()

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        filename = _safe_name(project.title, "docx")
        return Response(
            content=buf.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erreur DOCX : {str(e)}")


# ── Markdown (.md) ─────────────────────────────────────────────────────────────

@router.post("/markdown")
async def export_markdown(body: dict, db: AsyncSession = Depends(get_db)):
    project_id = body.get("project_id")
    if not project_id:
        raise HTTPException(status_code=400, detail="project_id requis")
    project, chapters = await _get_project_chapters(project_id, db)

    def html_to_md(html):
        h = re.sub(r'<strong>(.*?)</strong>', r'**\1**', html)
        h = re.sub(r'<em>(.*?)</em>', r'*\1*', h)
        h = re.sub(r'<h1>(.*?)</h1>', r'# \1', h)
        h = re.sub(r'<h2>(.*?)</h2>', r'## \1', h)
        h = re.sub(r'<h3>(.*?)</h3>', r'### \1', h)
        h = re.sub(r'<br\s*/?>', '\n', h)
        h = re.sub(r'</p>', '\n\n', h)
        h = re.sub(r'<[^>]+>', '', h)
        return h.strip()

    lines = [f"# {project.title}\n"]
    if project.subtitle: lines.append(f"*{project.subtitle}*\n")
    if project.author: lines.append(f"**{project.author}**\n")
    if project.synopsis: lines.append(f"\n> {project.synopsis}\n")
    lines.append("\n---\n")

    for i, chapter in enumerate(chapters):
        lines.append(f"\n## Chapitre {i+1} — {chapter.title}\n")
        if chapter.subtitle: lines.append(f"*{chapter.subtitle}*\n")
        if chapter.pov: lines.append(f"*PDV : {chapter.pov}*\n")
        lines.append("\n" + html_to_md(chapter.content or "") + "\n")

    content = "\n".join(lines)
    filename = _safe_name(project.title, "md")
    return Response(
        content=content.encode("utf-8"),
        media_type="text/markdown; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


# ── Texte brut (.txt) ──────────────────────────────────────────────────────────

@router.post("/txt")
async def export_txt(body: dict, db: AsyncSession = Depends(get_db)):
    project_id = body.get("project_id")
    if not project_id:
        raise HTTPException(status_code=400, detail="project_id requis")
    project, chapters = await _get_project_chapters(project_id, db)

    def html_to_text(html):
        h = re.sub(r'<br\s*/?>', '\n', html)
        h = re.sub(r'</p>', '\n\n', h)
        h = re.sub(r'<[^>]+>', '', h)
        return h.strip()

    lines = [project.title.upper(), "=" * len(project.title)]
    if project.subtitle: lines.append(project.subtitle)
    if project.author: lines.append(f"par {project.author}")
    lines.extend(["", ""])

    for i, chapter in enumerate(chapters):
        lines.append(f"CHAPITRE {i+1} — {chapter.title.upper()}")
        lines.append("-" * 40)
        if chapter.subtitle: lines.append(f"[{chapter.subtitle}]")
        lines.extend(["", html_to_text(chapter.content or ""), "", ""])

    content = "\n".join(lines)
    filename = _safe_name(project.title, "txt")
    return Response(
        content=content.encode("utf-8"),
        media_type="text/plain; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )
