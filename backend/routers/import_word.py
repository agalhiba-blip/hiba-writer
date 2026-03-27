from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from backend.database import get_db
from backend.models import Chapter, Project
from typing import Optional
import re
import io

router = APIRouter()

# Styles Word reconnus comme titres de chapitre (FR + EN + variantes)
HEADING_STYLES = {
    'heading 1', 'heading 2', 'heading1', 'heading2',
    'titre 1', 'titre 2', 'titre1', 'titre2',
    'ttre 1', 'ttre 2',   # typos courantes
    'heading', 'titre', 'chapitre', 'chapter',
    'h1', 'h2',
}

# Patterns de texte détectés comme titres même sans style Word
CHAPTER_PATTERNS = re.compile(
    r'^(chapitre|chapter|partie|part|prologue|épilogue|epilogue|acte|tome|volume)\s*[\d\w]*',
    re.IGNORECASE,
)


def is_heading_para(para, mode: str) -> bool:
    """Détermine si un paragraphe est un titre de chapitre."""
    if mode != 'auto':
        return False
    style_name = (para.style.name or '').lower().strip() if para.style else ''
    # Vérifier le style Word
    for h in HEADING_STYLES:
        if h in style_name:
            return True
    text = para.text.strip()
    if not text:
        return False
    # Vérifier le pattern textuel
    if CHAPTER_PATTERNS.match(text):
        return True
    # Ligne courte tout en majuscules (ex: "PREMIER ACTE")
    if len(text) <= 60 and text == text.upper() and len(text.split()) >= 2:
        return True
    # Paragraphe court entièrement en gras (titre sans style)
    if para.runs and all(r.bold for r in para.runs if r.text.strip()):
        if len(text) <= 80:
            return True
    return False


def para_to_html(para) -> str:
    """Convertit un paragraphe Word en HTML avec mise en forme."""
    html = '<p>'
    for run in para.runs:
        t = run.text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        if not t:
            continue
        if run.bold and run.italic:
            t = f'<strong><em>{t}</em></strong>'
        elif run.bold:
            t = f'<strong>{t}</strong>'
        elif run.italic:
            t = f'<em>{t}</em>'
        html += t
    html += '</p>'
    return html


def count_words(html_content: str) -> int:
    text = re.sub(r'<[^>]+>', ' ', html_content or '')
    return len(text.split())


@router.post("")
async def import_word(
    file: UploadFile = File(...),
    project_id: int = Form(0),
    mode: str = Form("auto"),
    project_title: str = Form(""),   # Nouveau : titre du roman à créer si besoin
    db: AsyncSession = Depends(get_db),
):
    """
    Importe un fichier .docx et crée des chapitres.
    Crée le projet dans la même requête si project_id == 0 ou introuvable.
    """
    if not file.filename.lower().endswith(('.docx', '.doc')):
        raise HTTPException(status_code=400, detail="Seuls les fichiers .docx sont supportés côté serveur")

    # ── Récupérer ou créer le projet dans la MÊME requête (évite le bug Vercel) ──
    project = None
    if project_id:
        project = await db.get(Project, project_id)

    if not project:
        title = project_title.strip() or file.filename.replace('.docx', '').replace('.doc', '')
        project = Project(title=title[:255])
        db.add(project)
        await db.flush()   # Obtenir l'ID sans commit définitif

    try:
        import docx as docx_lib
    except ImportError:
        raise HTTPException(
            status_code=501,
            detail="python-docx non disponible. Utilisez l'import navigateur (mammoth.js)."
        )

    content = await file.read()
    try:
        doc = docx_lib.Document(io.BytesIO(content))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Impossible de lire le fichier : {e}")

    base_name = file.filename.replace('.docx', '').replace('.doc', '')
    chapters_data = []
    current_title = base_name
    current_paras: list[str] = []

    for para in doc.paragraphs:
        if is_heading_para(para, mode):
            if current_paras:
                chapters_data.append({'title': current_title, 'content': ''.join(current_paras)})
            current_title = para.text.strip() or f"Chapitre {len(chapters_data) + 1}"
            current_paras = []
        else:
            text = para.text.strip()
            if text:
                current_paras.append(para_to_html(para))

    if current_paras:
        chapters_data.append({'title': current_title, 'content': ''.join(current_paras)})

    # Si aucun chapitre trouvé, mettre tout le document dans un seul chapitre
    if not chapters_data:
        full_html = ''.join(
            para_to_html(p) for p in doc.paragraphs if p.text.strip()
        )
        if not full_html:
            raise HTTPException(status_code=400, detail="Aucun contenu trouvé dans le fichier")
        chapters_data = [{'title': base_name, 'content': full_html}]

    # Calculer le prochain order_index
    result = await db.execute(
        select(Chapter).where(Chapter.project_id == project.id).order_by(Chapter.order_index.desc())
    )
    last = result.scalars().first()
    next_index = (last.order_index + 1) if last else 0

    created_titles = []
    for i, ch_data in enumerate(chapters_data):
        html_content = ch_data['content']
        chapter = Chapter(
            project_id=project.id,
            title=ch_data['title'][:255],
            content=html_content,
            order_index=next_index + i,
            word_count=count_words(html_content),
        )
        db.add(chapter)
        created_titles.append(ch_data['title'])

    await db.commit()
    return {
        "project_id": project.id,
        "created": len(created_titles),
        "chapters": created_titles,
    }
